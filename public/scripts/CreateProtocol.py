import os
import sys
import json
import base64
import io
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.shared import Inches

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

base64_data = sys.argv[1]
json_data = base64.b64decode(base64_data).decode('utf-8')
worksDB = json.loads(json_data)

doc = Document()
style = doc.styles['Normal']
doc_font = style.font
doc_font.name = 'Times New Roman'
doc_font.size = Pt(14)

# Добавляем заголовок документа
headline = doc.add_paragraph('Заключение заседания Методического совета ИРНИТУ')
headline.alignment = 1
doc.add_paragraph('от ____ _______________ ______ г. Протокол №____')
doc.add_paragraph('рассмотрение печатных изданий, которые по условиям методики формирования и распределения стимулирующих выплат подлежат включению в Рейтинг')

# Устанавливаем альбомную ориентацию для каждого раздела
for section in doc.sections:
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height
    section.left_margin = Inches(0.2)  # Левый отступ
    section.right_margin = Inches(0.2)  # Правый отступ
    section.top_margin = Inches(0.6)  # Верхний отступ
    section.bottom_margin = Inches(0.4)  # Нижний отступ

table_text_style = doc.styles.add_style('TableTextStyle', WD_STYLE_TYPE.PARAGRAPH)
table_text_font = table_text_style.font
table_text_font.name = 'Times New Roman'
table_text_font.size = Pt(10)  # Размер шрифта для текста в таблице

# Добавляем таблицу
table = doc.add_table(rows=1, cols=10)
table.style = 'Table Grid'

# Создаем стиль для заголовков столбцов
header_style = doc.styles.add_style('HeaderStyle', WD_STYLE_TYPE.PARAGRAPH)
header_font = header_style.font
header_font.name = 'Times New Roman'
header_font.size = Pt(10)
header_p = header_style.paragraph_format
header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Устанавливаем заголовки столбцов
headers = ['№', 'Вид публикации', 'Гриф/без грифа', 'Наименование', 'Авторы (ФИО полностью)', 'Институт (факультет)', 'Итоговый балл', 'Кол-во печатных листов', 'Издатель', 'Год издания']
for i, header in enumerate(headers):
    # Проверяем, есть ли уже абзацы в ячейке, иначе добавляем новый
    if len(table.cell(0, i).paragraphs) == 0:
        table.cell(0, i).add_paragraph(header)
    else:
        table.cell(0, i).paragraphs[0].text = header
    # Устанавливаем стиль для абзаца
    table.cell(0, i).paragraphs[0].style = header_style
    table.cell(0, i).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Добавляем данные в таблицу
data = []
i = 1

for work in worksDB:
    type = work.get('type')
    stamp = work.get('stamp')
    name_work = work.get('name_work')
    autor = work.get('autor')
    final_grade = work.get('final_grade')
    pages_number = work.get('pages_number')
    publisher = work.get('publisher')
    if (publisher is None):
        publisher = ''
    publishing_year = work.get('publishing_year')
    if (publishing_year is None):
        publishing_year = ''
    print(publisher)
    data.append((i, type, stamp, name_work, autor, '', final_grade, pages_number, publisher, publishing_year))
    i = i + 1

for row_data in data:
    cells = table.add_row().cells
    for i, cell_value in enumerate(row_data):
        cells[i].text = str(cell_value)
        # Устанавливаем стиль для абзаца в ячейке
        cells[i].paragraphs[0].style = header_style
        cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Текущая директория
current_dir = os.path.dirname(os.path.abspath(__file__))
# Поднимаемся на уровень выше
public_dir = os.path.abspath(os.path.join(current_dir, '..'))
# Переход в нужную директорию + название файла
save_path = os.path.join(public_dir, 'protocolFile', 'Протокол.docx')
# Сохранение фалйа
doc.save(save_path)
